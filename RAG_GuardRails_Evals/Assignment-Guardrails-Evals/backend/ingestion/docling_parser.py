"""
docling_parser.py - Parse documents using Docling with hierarchical structure extraction
"""
from __future__ import annotations
import os
import uuid
import logging
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


def parse_document(file_path: str) -> list[dict[str, Any]]:
    """
    Parse a document using Docling and extract hierarchical chunks with metadata.

    Returns a list of chunk dicts with fields:
        chunk_id, parent_chunk_id, text, chunk_type, section_title,
        page_number, source_document, collection, access_roles, level
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    # Route to appropriate parser based on file type
    if suffix == ".pdf":
        return _parse_with_docling(file_path)
    elif suffix in (".docx", ".doc"):
        return _parse_docx(file_path)
    elif suffix == ".md":
        return _parse_markdown(file_path)
    elif suffix == ".csv":
        return _parse_csv(file_path)
    else:
        logger.warning(f"Unsupported file type: {suffix}, attempting Docling fallback")
        return _parse_with_docling(file_path)


def _parse_with_docling(file_path: str) -> list[dict[str, Any]]:
    """Use Docling to parse PDF files with structural awareness."""
    try:
        from docling.document_converter import DocumentConverter
        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling_core.types.doc import DoclingDocument

        converter = DocumentConverter()
        result = converter.convert(file_path)
        doc = result.document

        chunks = []
        doc_id = str(uuid.uuid4())
        current_section = "Introduction"
        current_section_id = None
        page_number = 1

        # Create a root/document-level chunk
        doc_chunk_id = str(uuid.uuid4())
        chunks.append({
            "chunk_id": doc_chunk_id,
            "parent_chunk_id": None,
            "text": f"Document: {Path(file_path).name}",
            "chunk_type": "heading",
            "section_title": "Document Root",
            "page_number": 1,
            "source_document": Path(file_path).name,
            "collection": None,  # Will be set by ingest.py
            "access_roles": None,  # Will be set by ingest.py
            "level": 0,
        })

        # Extract content items from the Docling document
        for item, level in doc.iterate_items():
            item_type = type(item).__name__

            # Try to get page number
            try:
                if hasattr(item, 'prov') and item.prov:
                    page_number = item.prov[0].page_no if item.prov else page_number
            except Exception:
                pass

            # Determine chunk type and extract text
            text = ""
            chunk_type = "text"
            parent_id = current_section_id if current_section_id else doc_chunk_id

            if "SectionHeader" in item_type or "Heading" in item_type:
                try:
                    text = item.text if hasattr(item, 'text') else str(item)
                except Exception:
                    text = str(item)
                chunk_type = "heading"
                current_section = text
                section_chunk_id = str(uuid.uuid4())
                current_section_id = section_chunk_id
                parent_id = doc_chunk_id

                chunks.append({
                    "chunk_id": section_chunk_id,
                    "parent_chunk_id": parent_id,
                    "text": text,
                    "chunk_type": chunk_type,
                    "section_title": text,
                    "page_number": page_number,
                    "source_document": Path(file_path).name,
                    "collection": None,
                    "access_roles": None,
                    "level": level if level else 1,
                })
                continue

            elif "Table" in item_type:
                try:
                    text = item.export_to_markdown() if hasattr(item, 'export_to_markdown') else str(item)
                except Exception:
                    try:
                        text = item.text if hasattr(item, 'text') else str(item)
                    except Exception:
                        text = "[Table content]"
                chunk_type = "table"

            elif "Code" in item_type or "CodeItem" in item_type:
                try:
                    text = item.text if hasattr(item, 'text') else str(item)
                except Exception:
                    text = str(item)
                chunk_type = "code"

            else:
                # Text paragraph
                try:
                    text = item.text if hasattr(item, 'text') else str(item)
                except Exception:
                    text = str(item)
                chunk_type = "text"

            if not text or len(text.strip()) < 10:
                continue

            chunks.append({
                "chunk_id": str(uuid.uuid4()),
                "parent_chunk_id": parent_id,
                "text": text.strip(),
                "chunk_type": chunk_type,
                "section_title": current_section,
                "page_number": page_number,
                "source_document": Path(file_path).name,
                "collection": None,
                "access_roles": None,
                "level": (level if level else 2),
            })

        logger.info(f"Docling parsed {len(chunks)} chunks from {Path(file_path).name}")
        return chunks

    except Exception as e:
        logger.error(f"Docling parsing failed for {file_path}: {e}")
        # Fallback to basic text extraction
        return _fallback_text_parse(file_path)


def _parse_docx(file_path: str) -> list[dict[str, Any]]:
    """Parse .docx files using python-docx with hierarchical structure."""
    try:
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn

        doc = DocxDocument(file_path)
        chunks = []
        doc_chunk_id = str(uuid.uuid4())

        chunks.append({
            "chunk_id": doc_chunk_id,
            "parent_chunk_id": None,
            "text": f"Document: {Path(file_path).name}",
            "chunk_type": "heading",
            "section_title": "Document Root",
            "page_number": 1,
            "source_document": Path(file_path).name,
            "collection": None,
            "access_roles": None,
            "level": 0,
        })

        current_section = "Introduction"
        current_section_id = doc_chunk_id
        page_number = 1
        buffer_texts = []

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()

            if not text:
                continue

            if style_name.startswith("Heading"):
                # Flush buffer
                if buffer_texts:
                    combined = " ".join(buffer_texts)
                    if len(combined.strip()) >= 10:
                        chunks.append({
                            "chunk_id": str(uuid.uuid4()),
                            "parent_chunk_id": current_section_id,
                            "text": combined,
                            "chunk_type": "text",
                            "section_title": current_section,
                            "page_number": page_number,
                            "source_document": Path(file_path).name,
                            "collection": None,
                            "access_roles": None,
                            "level": 2,
                        })
                    buffer_texts = []

                current_section = text
                section_id = str(uuid.uuid4())

                try:
                    level = int(style_name.replace("Heading ", ""))
                except ValueError:
                    level = 1

                parent_id = doc_chunk_id if level == 1 else current_section_id
                current_section_id = section_id

                chunks.append({
                    "chunk_id": section_id,
                    "parent_chunk_id": parent_id,
                    "text": text,
                    "chunk_type": "heading",
                    "section_title": text,
                    "page_number": page_number,
                    "source_document": Path(file_path).name,
                    "collection": None,
                    "access_roles": None,
                    "level": level,
                })
            else:
                buffer_texts.append(text)
                # Chunk at ~500 words
                if sum(len(t.split()) for t in buffer_texts) > 500:
                    combined = " ".join(buffer_texts)
                    chunks.append({
                        "chunk_id": str(uuid.uuid4()),
                        "parent_chunk_id": current_section_id,
                        "text": combined,
                        "chunk_type": "text",
                        "section_title": current_section,
                        "page_number": page_number,
                        "source_document": Path(file_path).name,
                        "collection": None,
                        "access_roles": None,
                        "level": 2,
                    })
                    buffer_texts = []

        # Flush remaining
        if buffer_texts:
            combined = " ".join(buffer_texts)
            if len(combined.strip()) >= 10:
                chunks.append({
                    "chunk_id": str(uuid.uuid4()),
                    "parent_chunk_id": current_section_id,
                    "text": combined,
                    "chunk_type": "text",
                    "section_title": current_section,
                    "page_number": page_number,
                    "source_document": Path(file_path).name,
                    "collection": None,
                    "access_roles": None,
                    "level": 2,
                })

        # Parse tables
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_text = "\n".join(rows)
            if len(table_text.strip()) >= 10:
                chunks.append({
                    "chunk_id": str(uuid.uuid4()),
                    "parent_chunk_id": current_section_id,
                    "text": f"[Table {i+1}]\n{table_text}",
                    "chunk_type": "table",
                    "section_title": current_section,
                    "page_number": page_number,
                    "source_document": Path(file_path).name,
                    "collection": None,
                    "access_roles": None,
                    "level": 2,
                })

        logger.info(f"python-docx parsed {len(chunks)} chunks from {Path(file_path).name}")
        return chunks

    except Exception as e:
        logger.error(f"DOCX parsing failed for {file_path}: {e}")
        return _fallback_text_parse(file_path)


def _parse_markdown(file_path: str) -> list[dict[str, Any]]:
    """Parse Markdown files preserving heading hierarchy."""
    chunks = []
    doc_chunk_id = str(uuid.uuid4())

    chunks.append({
        "chunk_id": doc_chunk_id,
        "parent_chunk_id": None,
        "text": f"Document: {Path(file_path).name}",
        "chunk_type": "heading",
        "section_title": "Document Root",
        "page_number": 1,
        "source_document": Path(file_path).name,
        "collection": None,
        "access_roles": None,
        "level": 0,
    })

    try:
        content = Path(file_path).read_text(encoding="utf-8")
    except Exception:
        content = Path(file_path).read_text(encoding="latin-1")

    lines = content.split("\n")
    current_section = "Introduction"
    current_section_id = doc_chunk_id
    section_stack: list[tuple[int, str, str]] = []  # (level, title, chunk_id)
    buffer: list[str] = []
    in_code_block = False
    code_buffer: list[str] = []

    # Approx page tracking (every ~50 lines = 1 page)
    page_number = 1

    def flush_buffer():
        nonlocal buffer
        text = "\n".join(buffer).strip()
        if len(text) >= 10:
            chunks.append({
                "chunk_id": str(uuid.uuid4()),
                "parent_chunk_id": current_section_id,
                "text": text,
                "chunk_type": "text",
                "section_title": current_section,
                "page_number": page_number,
                "source_document": Path(file_path).name,
                "collection": None,
                "access_roles": None,
                "level": len(section_stack) + 1,
            })
        buffer = []

    for line_num, line in enumerate(lines):
        if line_num > 0 and line_num % 50 == 0:
            page_number += 1

        # Code blocks
        if line.startswith("```"):
            if in_code_block:
                in_code_block = False
                code_text = "\n".join(code_buffer).strip()
                if code_text:
                    chunks.append({
                        "chunk_id": str(uuid.uuid4()),
                        "parent_chunk_id": current_section_id,
                        "text": f"```\n{code_text}\n```",
                        "chunk_type": "code",
                        "section_title": current_section,
                        "page_number": page_number,
                        "source_document": Path(file_path).name,
                        "collection": None,
                        "access_roles": None,
                        "level": len(section_stack) + 1,
                    })
                code_buffer = []
            else:
                flush_buffer()
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        # Headings
        heading_match = None
        for prefix_len in range(6, 0, -1):
            prefix = "#" * prefix_len + " "
            if line.startswith(prefix):
                heading_match = (prefix_len, line[len(prefix):].strip())
                break

        if heading_match:
            flush_buffer()
            h_level, h_text = heading_match
            current_section = h_text
            section_id = str(uuid.uuid4())

            # Find parent
            parent_id = doc_chunk_id
            for s_level, s_title, s_id in reversed(section_stack):
                if s_level < h_level:
                    parent_id = s_id
                    break

            # Update stack
            section_stack = [(l, t, i) for l, t, i in section_stack if l < h_level]
            section_stack.append((h_level, h_text, section_id))
            current_section_id = section_id

            chunks.append({
                "chunk_id": section_id,
                "parent_chunk_id": parent_id,
                "text": h_text,
                "chunk_type": "heading",
                "section_title": h_text,
                "page_number": page_number,
                "source_document": Path(file_path).name,
                "collection": None,
                "access_roles": None,
                "level": h_level,
            })
        else:
            stripped = line.strip()
            if stripped:
                buffer.append(stripped)
                # Chunk at ~400 words
                if sum(len(l.split()) for l in buffer) > 400:
                    flush_buffer()

    flush_buffer()
    logger.info(f"Markdown parsed {len(chunks)} chunks from {Path(file_path).name}")
    return chunks


def _parse_csv(file_path: str) -> list[dict[str, Any]]:
    """Parse CSV files - each row becomes a chunk (for HR data)."""
    import pandas as pd

    chunks = []
    doc_chunk_id = str(uuid.uuid4())

    chunks.append({
        "chunk_id": doc_chunk_id,
        "parent_chunk_id": None,
        "text": f"Document: {Path(file_path).name} - Employee Records",
        "chunk_type": "heading",
        "section_title": "HR Employee Data",
        "page_number": 1,
        "source_document": Path(file_path).name,
        "collection": None,
        "access_roles": None,
        "level": 0,
    })

    try:
        df = pd.read_csv(file_path)

        # Group rows into batches of 5 to avoid too many tiny chunks
        batch_size = 5
        for i in range(0, len(df), batch_size):
            batch = df.iloc[i:i+batch_size]
            text_parts = []
            for _, row in batch.iterrows():
                row_text = " | ".join([f"{col}: {val}" for col, val in row.items() if str(val) != "nan"])
                text_parts.append(row_text)

            text = "\n".join(text_parts)
            chunks.append({
                "chunk_id": str(uuid.uuid4()),
                "parent_chunk_id": doc_chunk_id,
                "text": text,
                "chunk_type": "table",
                "section_title": "Employee Records",
                "page_number": (i // 50) + 1,
                "source_document": Path(file_path).name,
                "collection": None,
                "access_roles": None,
                "level": 1,
            })

        logger.info(f"CSV parsed {len(chunks)} chunks from {Path(file_path).name}")
    except Exception as e:
        logger.error(f"CSV parsing failed for {file_path}: {e}")

    return chunks


def _fallback_text_parse(file_path: str) -> list[dict[str, Any]]:
    """Basic fallback parser that reads raw text."""
    chunks = []
    doc_chunk_id = str(uuid.uuid4())

    try:
        # Try PyMuPDF for PDFs
        if file_path.lower().endswith(".pdf"):
            import fitz
            doc = fitz.open(file_path)
            full_text = []
            for page_num, page in enumerate(doc, 1):
                text = page.get_text()
                if text.strip():
                    full_text.append((page_num, text))

            chunks.append({
                "chunk_id": doc_chunk_id,
                "parent_chunk_id": None,
                "text": f"Document: {Path(file_path).name}",
                "chunk_type": "heading",
                "section_title": "Document Root",
                "page_number": 1,
                "source_document": Path(file_path).name,
                "collection": None,
                "access_roles": None,
                "level": 0,
            })

            for page_num, text in full_text:
                # Chunk each page into ~500-word segments
                words = text.split()
                for i in range(0, len(words), 500):
                    segment = " ".join(words[i:i+500]).strip()
                    if len(segment) >= 10:
                        chunks.append({
                            "chunk_id": str(uuid.uuid4()),
                            "parent_chunk_id": doc_chunk_id,
                            "text": segment,
                            "chunk_type": "text",
                            "section_title": f"Page {page_num}",
                            "page_number": page_num,
                            "source_document": Path(file_path).name,
                            "collection": None,
                            "access_roles": None,
                            "level": 1,
                        })
        else:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            chunks.append({
                "chunk_id": doc_chunk_id,
                "parent_chunk_id": None,
                "text": f"Document: {Path(file_path).name}\n\n{content[:5000]}",
                "chunk_type": "text",
                "section_title": "Document Root",
                "page_number": 1,
                "source_document": Path(file_path).name,
                "collection": None,
                "access_roles": None,
                "level": 0,
            })
    except Exception as e:
        logger.error(f"Fallback parsing failed for {file_path}: {e}")
        chunks.append({
            "chunk_id": doc_chunk_id,
            "parent_chunk_id": None,
            "text": f"Could not parse {Path(file_path).name}",
            "chunk_type": "text",
            "section_title": "Parse Error",
            "page_number": 1,
            "source_document": Path(file_path).name,
            "collection": None,
            "access_roles": None,
            "level": 0,
        })

    return chunks

